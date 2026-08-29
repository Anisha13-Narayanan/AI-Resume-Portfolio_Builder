from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class ResumeGenerator:

    def __init__(self):
        self.document = Document()

        # Page margins
        section = self.document.sections[0]

        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

        self._set_default_font()

    # ---------------------------------------------------------
    # DEFAULT FONT
    # ---------------------------------------------------------

    def _set_default_font(self):

        styles = self.document.styles

        normal = styles["Normal"]

        normal.font.name = "Arial"
        normal.font.size = Pt(10)

        normal._element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "Arial"
        )

    # ---------------------------------------------------------
    # ADD SECTION HEADING
    # ---------------------------------------------------------

    def _add_heading(self, text):

        paragraph = self.document.add_paragraph()

        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(3)

        run = paragraph.add_run(
            text.upper()
        )

        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(11)

        # Bottom border
        p = paragraph._p
        pPr = p.get_or_add_pPr()

        pBdr = OxmlElement("w:pBdr")

        bottom = OxmlElement("w:bottom")

        bottom.set(
            qn("w:val"),
            "single"
        )

        bottom.set(
            qn("w:sz"),
            "6"
        )

        bottom.set(
            qn("w:space"),
            "1"
        )

        pBdr.append(bottom)

        pPr.append(pBdr)

        return paragraph

    # ---------------------------------------------------------
    # HEADER
    # ---------------------------------------------------------

    def add_header(
        self,
        name,
        email="",
        phone=""
    ):

        paragraph = self.document.add_paragraph()

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        paragraph.paragraph_format.space_after = Pt(2)

        run = paragraph.add_run(
            name
        )

        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(20)

        contact_parts = []

        if email:
            contact_parts.append(email)

        if phone:
            contact_parts.append(phone)

        if contact_parts:

            contact = self.document.add_paragraph()

            contact.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            contact.paragraph_format.space_after = Pt(5)

            run = contact.add_run(
                " | ".join(contact_parts)
            )

            run.font.name = "Arial"
            run.font.size = Pt(9)

    # ---------------------------------------------------------
    # SUMMARY
    # ---------------------------------------------------------

    def add_summary(self, summary):

        if not summary:
            return

        self._add_heading(
            "Professional Summary"
        )

        paragraph = self.document.add_paragraph()

        paragraph.paragraph_format.space_after = Pt(3)

        paragraph.add_run(
            summary
        )

    # ---------------------------------------------------------
    # SKILLS
    # ---------------------------------------------------------

    def add_skills(self, skills):

        if not skills:
            return

        self._add_heading(
            "Technical Skills"
        )

        paragraph = self.document.add_paragraph()

        paragraph.paragraph_format.space_after = Pt(3)

        if isinstance(skills, list):

            skill_text = " • ".join(
                str(skill)
                for skill in skills
            )

        else:

            skill_text = str(skills)

        paragraph.add_run(
            skill_text
        )

    # ---------------------------------------------------------
    # EDUCATION
    # ---------------------------------------------------------

    def add_education(self, education):

        if not education:
            return

        self._add_heading(
            "Education"
        )

        paragraph = self.document.add_paragraph()

        degree = education.get(
            "degree",
            ""
        )

        field = education.get(
            "field",
            ""
        )

        institution = education.get(
            "institution",
            ""
        )

        year = education.get(
            "graduation_year",
            ""
        )

        if degree:

            run = paragraph.add_run(
                degree
            )

            run.bold = True

        if field:

            paragraph.add_run(
                f" — {field}"
            )

        if institution:

            paragraph.add_run(
                f"\n{institution}"
            )

        if year:

            paragraph.add_run(
                f" | {year}"
            )

        paragraph.paragraph_format.space_after = Pt(3)

    # ---------------------------------------------------------
    # PROJECTS
    # ---------------------------------------------------------

    def add_projects(self, projects):

        if not projects:
            return

        self._add_heading(
            "Projects"
        )

        for project in projects:

            name = project.get(
                "name",
                ""
            )

            description = project.get(
                "description",
                ""
            )

            technologies = project.get(
                "technologies",
                []
            )

            paragraph = self.document.add_paragraph()

            paragraph.paragraph_format.space_after = Pt(1)

            if name:

                run = paragraph.add_run(
                    name
                )

                run.bold = True
                run.font.size = Pt(10)

            if technologies:

                if isinstance(
                    technologies,
                    list
                ):

                    tech_text = ", ".join(
                        str(x)
                        for x in technologies
                    )

                else:

                    tech_text = str(
                        technologies
                    )

                paragraph.add_run(
                    f" | {tech_text}"
                )

            if description:

                desc = self.document.add_paragraph(
                    style=None
                )

                desc.paragraph_format.left_indent = (
                    Inches(0.15)
                )

                desc.paragraph_format.space_after = Pt(3)

                desc.add_run(
                    description
                )

    # ---------------------------------------------------------
    # EXPERIENCE
    # ---------------------------------------------------------

    def add_experience(self, experience):

        if not experience:
            return

        self._add_heading(
            "Experience"
        )

        paragraph = self.document.add_paragraph()

        paragraph.add_run(
            str(experience)
        )

    # ---------------------------------------------------------
    # GENERATE
    # ---------------------------------------------------------

    def generate(
        self,
        resume,
        output_path
    ):

        self.add_header(

            name=resume.get(
                "name",
                ""
            ),

            email=resume.get(
                "email",
                ""
            ),

            phone=resume.get(
                "phone",
                ""
            )
        )

        self.add_summary(
            resume.get(
                "summary",
                ""
            )
        )

        self.add_skills(
            resume.get(
                "skills",
                []
            )
        )

        self.add_education(
            resume.get(
                "education",
                {}
            )
        )

        self.add_projects(
            resume.get(
                "projects",
                []
            )
        )

        self.add_experience(
            resume.get(
                "experience",
                ""
            )
        )

        self.document.save(
            output_path
        )

        return output_path